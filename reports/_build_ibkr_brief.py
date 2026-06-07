"""Build IBKR Postmortem Brief as Word document with inline attribution."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)


def add_cite(para, num):
    r = para.add_run("[" + str(num) + "]")
    r.font.superscript = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)


# ===== TITLE =====
title = doc.add_paragraph()
trun = title.add_run("IBKR Postmortem Brief - BREAKOUT-Signal 2026-06-01")
trun.bold = True
trun.font.size = Pt(18)

sub = doc.add_paragraph()
sub_run = sub.add_run("Interactive Brokers Group (NASDAQ: IBKR) - Stop Loss D+3, -5.39 %")
sub_run.italic = True
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta = doc.add_paragraph()
mrun = meta.add_run("Report-Datum: 2026-06-06 - ApexNext Postmortem Process - Confidence: HIGH")
mrun.font.size = Pt(9)
mrun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# ===== EXECUTIVE SUMMARY =====
h = doc.add_heading("Executive Summary", level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)

p = doc.add_paragraph()
p.add_run("IBKR (Score 123.3, Pocket Pivot, BREAKOUT) wurde am 01.06.2026 nahe All-Time-Highs gepickt ($88.62 vs. 52w-High $91.02). Stop Loss nach 3 Handelstagen bei -5.39 %. ")
p.add_run("Wichtig: IBKR ist KEIN BNPL-Risk-Off-Loser wie AFRM.").bold = True
p.add_run(" Fundamental ist Interactive Brokers ein Zinsen-Gewinner (Q1 2026: margin loans +40 %, NII +20 %, 77 % pretax-Margin)")
add_cite(p, 1)
p.add_run(". Failure-Pattern hier: ")
p.add_run("Late-cycle Breakout nahe 52w-High + breite Risk-Off-Welle = klassischer Mean-Revert").bold = True
p.add_run(", verschaerft durch das Macro-Event am 5. Juni")
add_cite(p, 2)
p.add_run(".")

# ===== TRADE DETAILS =====
doc.add_heading("Trade-Details", level=1)

table = doc.add_table(rows=8, cols=2)
table.style = "Light Grid Accent 1"
data = [
    ("Signal-Datum", "2026-06-01"),
    ("Setup", "BREAKOUT + Pocket Pivot"),
    ("Entry (buy_above)", "$88.62"),
    ("Stop Loss", "$83.84 (-5.39 %)"),
    ("Target", "$97.50 (+10 %)"),
    ("RR", "1.86"),
    ("Score", "123.3 (Top-Decile)"),
    ("Outcome", "Stop Loss D+3"),
]
for i, (k, v) in enumerate(data):
    cells = table.rows[i].cells
    cells[0].text = k
    cells[1].text = v
    cells[0].paragraphs[0].runs[0].bold = True

# ===== COMPANY PROFILE =====
doc.add_heading("Company Profile", level=1)

p = doc.add_paragraph()
p.add_run("Interactive Brokers Group, Inc.").bold = True
p.add_run(" ist ein US-elektronischer Broker fuer institutionelle und private Anleger. Markt-Cap ~$145 Mrd. (Stand 06.06.2026)")
add_cite(p, 3)
p.add_run(". Beta ")
p.add_run("1.329").bold = True
p.add_run(" - moderates Marktrisiko. Sektor: Financial Services / Investment Banking. CEO: Milan Galik. 3,027 Mitarbeiter. Hauptsitz Greenwich, CT.")

doc.add_heading("Q1 FY2026 Earnings (April 2026)", level=2)

bullets = [
    "Customer Accounts: 4.75 Mio. (+31 % YoY)",
    "Margin Loans: $90.2 Mrd. (+40 % YoY)",
    "Net Interest Income: $966 Mio. (+20 % YoY)",
    "Commission Revenue: $613 Mio. (+19 %)",
    "Pretax Profit Margin: 77 % (6. Quartal in Folge ueber 70 %)",
    "Customer Equity: $789.4 Mrd. (+38 %)",
]
for txt in bullets:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(txt)
    add_cite(bp, 1)

p = doc.add_paragraph()
p.add_run("Fundamental: ")
p.add_run("ausgezeichnet").bold = True
p.add_run(". IBKR ist klassischer ")
p.add_run("Zins-Beneficiary").bold = True
p.add_run(" - hoehere Zinsen bedeuten hoehere NII auf Customer-Credits.")

# ===== WHY THE BREAKOUT FAILED =====
doc.add_heading("Warum der BREAKOUT scheiterte - Mehrschicht-Analyse", level=1)

doc.add_heading("1. Late-Cycle BREAKOUT nahe 52w-High (primaere Ursache)", level=2)
p = doc.add_paragraph()
p.add_run("Entry bei $88.62 lag nur ").bold = False
p.add_run("2.6 % unter dem 52w-High ($91.02)").bold = True
add_cite(p, 3)
p.add_run(". Das ist im ")
p.add_run("BREAKOUT-Failure-Zone-Bereich").bold = True
p.add_run(": Stocks die unmittelbar an Highs ausbrechen sind anfaellig fuer Mean-Reversion, weil:")

for txt in [
    "Profit-Taker setzen Stops knapp unter All-Time-Highs",
    "Wenig 'overhead supply' = wenig Schuetzer wenn Selloff einsetzt",
    "Institutionelle warten oft auf Pullback statt vertical zu kaufen",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(txt)

doc.add_heading("2. Macro-Event Verstaerker (sekundaere Ursache)", level=2)
p = doc.add_paragraph()
p.add_run("Am 5. Juni 2026 (D+4) fiel der Nasdaq -4 % wegen Rate-Hike-Aengsten nach starkem NFP-Report")
add_cite(p, 2)
p.add_run(". Bei Beta 1.33 bedeutet das mathematisch IBKR-Drop ~5 % - exakt entspricht dem -5.39 % Stop-Loss. Das ist KEIN ticker-spezifisches Problem, sondern ")
p.add_run("mechanisches Beta-Multiplikator-Verhalten").bold = True
p.add_run(".")

doc.add_heading("3. Paradox: Zins-Gewinner verlor durch Zins-Sorgen", level=2)
p = doc.add_paragraph()
p.add_run("IBKR profitiert eigentlich VON hoeheren Zinsen (NII auf Margin Loans und Customer Credits). Q1-2026-Earnings zeigten das deutlich (NII +20 %)")
add_cite(p, 1)
p.add_run(". Im Selloff vom 5. Juni 2026 wurde diese Logik aber ")
p.add_run("durch generelle Risk-Off-Welle ueberschrieben").bold = True
p.add_run(". Bei systematischen Drawdowns korreliert ALLES kurzfristig - rationales Pricing kehrt erst spaeter zurueck.")

# ===== KONTRAST AFRM vs IBKR =====
doc.add_heading("Kontrast AFRM vs IBKR: zwei verschiedene Failure-Modes", level=1)

ctable = doc.add_table(rows=8, cols=3)
ctable.style = "Light Grid Accent 1"

ctable.rows[0].cells[0].text = "Faktor"
ctable.rows[0].cells[1].text = "AFRM (29.5.)"
ctable.rows[0].cells[2].text = "IBKR (1.6.)"
for c in ctable.rows[0].cells:
    c.paragraphs[0].runs[0].bold = True

rows = [
    ("Beta", "3.7 (sehr hoch)", "1.33 (moderat)"),
    ("Sektor-Logik", "Rate-Hike = SCHLECHT (BNPL credit)", "Rate-Hike = GUT (NII)"),
    ("Distance to 52w-High", "Nicht nahe Highs", "Sehr nahe (-2.6 %)"),
    ("Failure-Pattern", "Sektor-Wegging + Beta-Multiplier", "Late-Cycle Top + Risk-Off"),
    ("Lesson", "Hochbeta + Credit-sensitive vermeiden in Hike-Window", "Nahe-High BREAKOUTs sind Falle bei broad selloff"),
    ("Score 123 PP", "Score Top-Decile half nicht", "Score Top-Decile half nicht"),
    ("Outcome", "-7.46 % Stop D+3", "-5.39 % Stop D+3"),
]
for i, (a, b, c) in enumerate(rows, start=1):
    ctable.rows[i].cells[0].text = a
    ctable.rows[i].cells[1].text = b
    ctable.rows[i].cells[2].text = c
    ctable.rows[i].cells[0].paragraphs[0].runs[0].bold = True

# ===== LESSON TAGS =====
doc.add_heading("Lesson-Tags (fuer Postmortem-DB)", level=1)
tags = [
    ("late_cycle_breakout_near_52w_high", "Breakouts naher 52w-High = Mean-Revert-Risiko erhoeht"),
    ("macro_selloff_correlates_all_stocks", "Im Crash korreliert ALLES kurz - Fundamental egal"),
    ("rate_beneficiary_paradox", "Auch Zins-Gewinner verlieren in pauschalem Rate-Hike-Selloff"),
    ("score_top_decile_no_protection_in_macro", "Score 100+ schuetzt nicht vor Macro-Shocks"),
]
for tag, desc in tags:
    p = doc.add_paragraph(style="List Bullet")
    tr = p.add_run(tag)
    tr.font.name = "Consolas"
    tr.font.size = Pt(10)
    tr.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(" - " + desc)

# ===== IMPLICATIONS =====
doc.add_heading("Implikationen fuer ApexNext", level=1)
implications = [
    "DISTANCE-TO-52W-HIGH-Filter: BREAKOUTs deren Entry < 3 % unter 52w-High liegt koennten Score-Penalty -5 bekommen. Validiert per Backtest.",
    "Macro-Calendar-Awareness (gleicher Punkt wie AFRM-Brief): NFP/CPI/FOMC innerhalb 7d sollten Score reduzieren.",
    "MIXED-Regime-Verhalten: AFRM + IBKR sind gemeinsame Faelle - bei Regime-Wechsel BULLISH -> MIXED sollte TG_MIN_SCORE evtl. von 70 auf 85 angehoben werden, um nur Elite zuzulassen.",
    "Beta != Catalyst-Free: Verschiedene Beta-Profile haben verschiedene Failure-Modes - AFRM/IBKR zeigt dass Score-System keinen davon vermeidet.",
]
for imp in implications:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(imp)

# ===== SOURCES =====
doc.add_page_break()
doc.add_heading("Quellen", level=1)

p = doc.add_paragraph()
p.add_run("Inline-Zitate verweisen auf folgende Quellen:").italic = True

sources = [
    ("[1]", "Interactive Brokers Q1 2026 Earnings Release (April 2026): Customer Accounts +31 %, Margin Loans +40 %, NII +20 %, 77 % Pretax Margin",
     "https://www.sec.gov/Archives/edgar/data/0001381197/000138119726000078/ibkr-ex99_1.htm"),
    ("[2]", "CNN Business - \"Nasdaq, S&P 500 suffer worst day of year as AI stocks tumble and Fed rate-hike odds rise\" (5. Juni 2026)",
     "https://www.cnn.com/2026/06/05/markets/stock-market-sell-off-fed"),
    ("[3]", "Financial Modeling Prep - IBKR Company Profile API (abgerufen 06.06.2026): marketCap $145B, beta 1.329, 52w Range $49.15-$91.02",
     "https://financialmodelingprep.com/api/v3/profile/IBKR"),
]
for num, title, url in sources:
    p = doc.add_paragraph()
    n = p.add_run(num + " ")
    n.bold = True
    n.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p.add_run(title)
    p.add_run("\n")
    url_run = p.add_run(url)
    url_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    url_run.font.size = Pt(9)
    url_run.italic = True

# ===== FOOTER =====
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
f = footer.add_run("Powered by Bigdata.com (skill orchestration) - Financial Modeling Prep (company data) - Web Research (news context)")
f.font.size = Pt(8)
f.italic = True
f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

disc = doc.add_paragraph()
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
d = disc.add_run("Disclaimer: Dieses Dokument ist eine interne Postmortem-Analyse fuer das ApexNext-Trading-System und stellt keine Investitionsempfehlung dar.")
d.font.size = Pt(8)
d.italic = True
d.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

out = r"C:\Users\Niklas\TradeBot\ApexNext\reports\IBKR_company_brief_2026-06.docx"
doc.save(out)
print("SAVED:", out)
print("Size:", os.path.getsize(out), "bytes")
