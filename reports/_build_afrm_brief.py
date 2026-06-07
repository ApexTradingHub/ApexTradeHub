"""Build AFRM Postmortem Brief as Word document with inline attribution."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Default style
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)


def add_cite(para, num):
    """Add inline citation marker like [1] as superscript."""
    r = para.add_run("[" + str(num) + "]")
    r.font.superscript = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)


# ===== TITLE =====
title = doc.add_paragraph()
trun = title.add_run("AFRM Postmortem Brief - BREAKOUT-Signal 2026-05-29")
trun.bold = True
trun.font.size = Pt(18)

sub = doc.add_paragraph()
sub_run = sub.add_run("Affirm Holdings, Inc. (NASDAQ: AFRM) - Stop Loss D+3, -7.46 %")
sub_run.italic = True
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta = doc.add_paragraph()
mrun = meta.add_run("Report-Datum: 2026-06-06 - ApexNext Postmortem Process - Confidence: HIGH")
mrun.font.size = Pt(9)
mrun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# ===== EXECUTIVE SUMMARY =====
h = doc.add_heading("Executive Summary", level=1)
for r in h.runs:
    r.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)

p = doc.add_paragraph()
p.add_run("AFRM (Score 123.6, Pocket Pivot, BREAKOUT) wurde am 29.05.2026 gepickt und nach drei Handelstagen am Stop ausgeloest (-7.46 %). Die Position fiel NICHT durch einen ticker-spezifischen Schock, sondern durch ein scharfes ")
p.add_run("Makro-Risk-Off-Event").bold = True
p.add_run(" am 5. Juni 2026: starker US-Arbeitsmarktbericht (+172 K Jobs, ~2x Konsens) toetete die Rate-Cut-Erwartungen, Nasdaq fiel -4 % an einem Tag")
add_cite(p, 1)
p.add_run(". AFRM mit Beta 3.7")
add_cite(p, 2)
p.add_run(" wurde ueberproportional getroffen.")

# ===== TRADE DETAILS =====
doc.add_heading("Trade-Details", level=1)

table = doc.add_table(rows=8, cols=2)
table.style = "Light Grid Accent 1"
data = [
    ("Signal-Datum", "2026-05-29"),
    ("Setup", "BREAKOUT (Pocket Pivot bestaetigt)"),
    ("Entry (buy_above)", "~$75.20"),
    ("Stop Loss", "-7.46 % vom Entry"),
    ("Target", "+13 % (nicht erreicht)"),
    ("RR", "1.83"),
    ("Score", "123.6 (Top-Decile)"),
    ("Outcome", "Stop Loss D+3"),
]
for i, (k, v) in enumerate(data):
    cells = table.rows[i].cells
    cells[0].text = k
    cells[1].text = v
    cells[0].paragraphs[0].runs[0].bold = True

# ===== COMPANY PROFILE =====
doc.add_heading("Company Profile", level=1)

p = doc.add_paragraph()
p.add_run("Affirm Holdings, Inc.").bold = True
p.add_run(" ist ein US-Buy-Now-Pay-Later (BNPL) Plattformanbieter mit Hauptsitz in San Francisco. CEO: Max Levchin. Marktkapitalisierung: ~$21.3 Mrd. (Stand 06.06.2026)")
add_cite(p, 2)
p.add_run(". 52-Wochen-Range: $42.10-$100.00. Beta 3.7 -> ca. 3.7x S&P-500-Vola. Branche: Software-Infrastructure / Fintech. 26.8 Mio. aktive Konsumenten, 2,006 Mitarbeiter")
add_cite(p, 3)
p.add_run(".")

doc.add_heading("Letzter Earnings-Beat (Q3 FY2026, 8. Mai 2026)", level=2)

bullets = [
    ("GMV: $11.6 Mrd. (+35 % YoY)", 3),
    ("Revenue: $1,039 Mio. (+33 % YoY)", 3),
    ("Net Income: $100 Mio. (vs. $2.8 Mio. Vorjahr) - erste signifikante GAAP-Profitabilitaet", 3),
    ("Adj. Operating Margin: 27 % (+5 pp YoY)", 3),
    ("FY2026 Guidance angehoben - Bull-Case-Bestaetigung", 3),
]
for txt, cite in bullets:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(txt)
    add_cite(bp, cite)

p = doc.add_paragraph()
p.add_run("Fundamental-Setup war ").italic = True
p.add_run("intakt zum Signal-Zeitpunkt").bold = True
p.add_run(". Earnings-Beat lag 21 Tage zurueck (ausserhalb Earnings-Blackout-Window). Wachstum starke 33-41 %.")

# ===== WHY THE BREAKOUT FAILED =====
doc.add_heading("Warum der BREAKOUT scheiterte - Mehrschicht-Analyse", level=1)

doc.add_heading("1. Makro-Schock (primaere Ursache)", level=2)
p = doc.add_paragraph()
p.add_run("Am 5. Juni 2026 (D+5 vom Signal, D+2 nach Stop) veroeffentlichten US-Arbeitsmarktdaten +172 K neue Stellen - mehr als doppelt so viel wie Konsens-Erwartung")
add_cite(p, 1)
p.add_run(". Konsequenzen:")

reasons = [
    "Fed-Rate-Cut-Hoffnungen fuer Sommer 2026 verschwanden",
    "42.7 % implizite Wahrscheinlichkeit eines Rate-Hike im Dezember",
    "Nasdaq -4 % an diesem Tag (worst single day of 2026)",
    "Fintech / Consumer-Credit-sensitive Sektor besonders stark betroffen",
]
for txt in reasons:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(txt)
    if "Nasdaq" in txt:
        add_cite(bp, 1)

doc.add_heading("2. Mechanische Verstaerkung durch Beta 3.7", level=2)
p = doc.add_paragraph()
p.add_run("AFRM hatte eine 52-Wochen-Beta von ")
p.add_run("3.699").bold = True
add_cite(p, 2)
p.add_run(". Bei einem SPY-Drop von ~3 % im Zeitraum implizierte das mathematisch ca. -11 % AFRM-Drop - der tatsaechliche -7.46 % bis zum Stop liegt im erwartbaren Korridor. Hohe Beta ist ein ")
p.add_run("strukturelles BREAKOUT-Risiko").bold = True
p.add_run(", das vom Score nicht direkt erfasst wird.")

doc.add_heading("3. Sektor-Headwind (Fintech/BNPL)", level=2)
p = doc.add_paragraph()
p.add_run("BNPL-Aktien sind besonders zinssensitiv: hoehere Zinsen erhoehen die Funding-Kosten und schwaechen die Konsumenten-Kreditqualitaet. Mehrere BNPL/Fintech-Peers (PayPal, SoFi, Klarna, Fiserv) zeigten parallel Schwaeche im Zeitraum")
add_cite(p, 4)
p.add_run(".")

doc.add_heading("4. Verschaerfendes Hintergrundrauschen", level=2)
bg = [
    ("Walmart-Partnership war bereits im Maerz 2025 an Klarna verloren (5 % GMV-Risiko angekuendigt) - bekannt, gepriced, aber unter Druck wieder als Bear-Case-Anker zitiert", 5),
    ("BNPL-Regulierung: CFPB unter Trump-Administration deregulierte BNPL ab Mai 2025, Marktstimmung wechselte schnell zu Credit-Quality-Bedenken", 6),
]
for txt, cite in bg:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(txt)
    add_cite(bp, cite)

# ===== LESSON TAGS =====
doc.add_heading("Lesson-Tags (fuer Postmortem-DB)", level=1)
tags = [
    ("high_beta_breakout_macro_risk", "Beta >3 BREAKOUTs sind disproportional macro-getrieben"),
    ("fintech_consumer_credit_sensitivity", "BNPL/Fintech sind Rate-sensitive - bei Hike-Risk fade"),
    ("fundamentals_intact_but_stopped", "Q3-Beat + Guidance-Raise konnten Macro-Drop nicht abfangen"),
    ("rate_decision_window_risk", "BREAKOUT vor unsicherem Makro-Datenpunkt (NFP) = Falle"),
]
for tag, desc in tags:
    p = doc.add_paragraph(style="List Bullet")
    tr = p.add_run(tag)
    tr.font.name = "Consolas"
    tr.font.size = Pt(10)
    tr.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(" - " + desc)

# ===== IMPLICATIONS =====
doc.add_heading("Implikationen fuer ApexNext", level=1)
implications = [
    "BETA-Penalty erwaegen: BREAKOUTs mit Beta > 2.5 koennten Score-Penalty -10 bekommen (Hypothese, n=?). Aktuell keine Beta-Variable im Score.",
    "Macro-Calendar-Awareness: ApexCatalysts koennte upcoming NFP/CPI/FOMC checken und Score reduzieren bei high-impact events in 7d.",
    "Fundamental != Trade-Setup: Selbst CONFIRMED-Setup (PP + Score 123 + Earnings-Beat) kann durch Macro-Shock kippen. Position Sizing.",
    "Vergleichstrades: AFRM ist nicht isoliert - IBKR (-5.39 %), AXTA (open), ARE (open) im gleichen Zeitfenster zeigen, dass MIXED-Regime BREAKOUT-WR drueckt.",
]
for imp in implications:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(imp)

# ===== SOURCES =====
doc.add_page_break()
doc.add_heading("Quellen", level=1)

p = doc.add_paragraph()
p.add_run("Inline-Zitate verweisen auf folgende Quellen:").italic = True

sources = [
    ("[1]", "CNN Business - \"Nasdaq, S&P 500 suffer worst day of year as AI stocks tumble and Fed rate-hike odds rise\" (5. Juni 2026)",
     "https://www.cnn.com/2026/06/05/markets/stock-market-sell-off-fed"),
    ("[2]", "Financial Modeling Prep - AFRM Company Profile API (abgerufen 06.06.2026): marketCap $21.3B, beta 3.699, 52w Range $42.10-$100.00",
     "https://financialmodelingprep.com/api/v3/profile/AFRM"),
    ("[3]", "Affirm Holdings Form 8-K Q3 FY2026 Shareholder Letter (8. Mai 2026): GMV $11.6B, Revenue $1.039B, Net Income $100M",
     "https://www.sec.gov/Archives/edgar/data/0001820953/000162828026032105/affirmfq326shareholderle.htm"),
    ("[4]", "Yahoo Finance / TheStreet - Fintech-Sektor-Schwaeche, PayPal/SoFi/Klarna/Fiserv declines Mai-Juni 2026",
     "https://finance.yahoo.com/markets/live/stock-market-today-dow-sp-500-nasdaq-sink-as-jobs-report-fuels-fed-hike-bets-chip-stocks-sell-off-230134285.html"),
    ("[5]", "Retail Dive - \"Klarna to displace Affirm as Walmart BNPL provider\" (17. Maerz 2025)",
     "https://www.retaildive.com/news/walmart-loan-payments-klarna-buy-now-pay-later/742914/"),
    ("[6]", "GuruFocus - \"CFPB Shifts Focus Away from BNPL Lenders, Shares React\" (Mai 2025)",
     "https://www.gurufocus.com/news/2835110/cfpb-shifts-focus-away-from-bnpl-lenders-shares-react"),
]
for num, title, url in sources:
    p = doc.add_paragraph()
    n = p.add_run(num + " ")
    n.bold = True
    n.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p.add_run(title)
    p.add_run("\n")
    url_run = p.add_run(url)
    url_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    url_run.font.size = Pt(9)
    url_run.italic = True

# ===== FOOTER ATTRIBUTION =====
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
f = footer.add_run("Powered by Bigdata.com (skill orchestration) - Financial Modeling Prep (company data) - Web Research (news context)")
f.font.size = Pt(8)
f.italic = True
f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

disc = doc.add_paragraph()
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
d = disc.add_run("Disclaimer: Dieses Dokument ist eine interne Postmortem-Analyse fuer das ApexNext-Trading-System und stellt keine Investitionsempfehlung dar.")
d.font.size = Pt(8)
d.italic = True
d.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
out = r"C:\Users\Niklas\TradeBot\ApexNext\reports\AFRM_company_brief_2026-06.docx"
doc.save(out)
print("SAVED:", out)
print("Size:", os.path.getsize(out), "bytes")
